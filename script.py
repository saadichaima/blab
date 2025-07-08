import os
import fitz  # pour les PDF
from docx import Document  # pour les fichiers Word
import openai
from dotenv import load_dotenv
from openai import AzureOpenAI
load_dotenv()

import os

# Configure OpenAI (Azure)
client = AzureOpenAI(
    api_key=os.getenv("AZURE_OPENAI_KEY"),
    api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
    azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT")
)

deployment_name = os.getenv("AZURE_OPENAI_DEPLOYMENT")



def extract_text_from_folder(folder_path):
    full_text = ""

    for file in os.listdir(folder_path):
        file_path = os.path.join(folder_path, file)

        if file.endswith(".pdf"):
            doc = fitz.open(file_path)
            for page in doc:
                full_text += page.get_text()
            doc.close()

        elif file.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8") as f:
                full_text += f.read()

        elif file.endswith(".docx"):
            doc = Document(file_path)
            for para in doc.paragraphs:
                full_text += para.text + "\n"

    return full_text
def chunk_text(text, max_length=1000, overlap=200):
    chunks = []
    start = 0
    while start < len(text):
        end = start + max_length
        chunk = text[start:end]
        chunks.append(chunk)
        start += max_length - overlap  # pour avoir un peu de recouvrement
    return chunks
from docx import Document
def generate_report_from_chunks(chunks):
    # On prend les chunks les plus pertinents ou les premiers (selon ton besoin)
    context = "\n\n".join(chunks[:5])  # Tu peux augmenter si besoin

    prompt = f"""
    Voici des extraits de documents clients :
    {context}

    Génère un rapport structuré comprenant les sections suivantes :
    1. Contexte et objectif de l'operation 
    2. Périmetre de l'étude marché ciblé
    3. Etats de produit commercialisés
    4. identifications des nouvelles performances visées
    5. synthèses des travaux 
    """
    response = client.chat.completions.create(
        model=deployment_name,  # modèle = nom du déploiement Azure
        messages=[
            {"role": "user", "content": prompt}
        ],
        max_tokens=1000,
        temperature=0.5
    )


    return response.choices[0].message.content

def save_report_to_word(report_text, output_path):
    doc = Document()
    doc.add_heading('Rapport Client Généré', 0)

    for line in report_text.split("\n"):
        if line.strip() == "":
            continue
        elif line.strip().endswith(":") or line.strip().startswith(("1.", "2.", "3.", "4.", "5.")):
            doc.add_heading(line.strip(), level=1)
        else:
            doc.add_paragraph(line.strip())

    doc.save(output_path)

if __name__ == "__main__":
    folder = "Doc"
    text = extract_text_from_folder(folder)
    chunks = chunk_text(text)

    print(f"Nombre de chunks générés : {len(chunks)}")
    print("\n--- Premier chunk ---\n")
    print(chunks[0])

    # Appel à l'IA
    report = generate_report_from_chunks(chunks)

    # Enregistrement du rapport
    save_report_to_word(report, "rapport_client.docx")
    print("📄 Rapport généré et enregistré dans 'rapport_client.docx'")
