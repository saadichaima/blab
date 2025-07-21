import os
import faiss
import numpy as np
import streamlit as st
import fitz
from docx import Document
from dotenv import load_dotenv
from openai import AzureOpenAI

# ─────── CONFIGURATION AZURE ────────
load_dotenv()
client = AzureOpenAI(
    api_key=os.getenv("AZURE_OPENAI_KEY"),
    azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
    api_version=os.getenv("AZURE_OPENAI_API_VERSION")
)
GPT_DEPLOYMENT = os.getenv("AZURE_OPENAI_DEPLOYMENT")
EMBEDDING_DEPLOYMENT = os.getenv("AZURE_OPENAI_EMBEDDING_DEPLOYMENT")

# ─────── FONCTIONS UTILES ───────────

def extract_text(file):
    if file.name.lower().endswith(".pdf"):
        doc = fitz.open(stream=file.read(), filetype="pdf")
        return "".join(page.get_text() for page in doc)
    elif file.name.lower().endswith(".docx"):
        document = Document(file)
        return "\n".join(para.text for para in document.paragraphs)
    return ""

def chunk_text(text, chunk_size=1000, overlap=200):
    words = text.split()
    return [" ".join(words[i:i + chunk_size])
            for i in range(0, len(words), chunk_size - overlap)]

def embed_texts(texts):
    response = client.embeddings.create(
        input=texts,
        model=EMBEDDING_DEPLOYMENT
    )
    return [np.array(e.embedding, dtype=np.float32) for e in response.data]

def build_faiss_index(chunks):
    vectors = embed_texts(chunks)
    dim = len(vectors[0])
    index = faiss.IndexFlatL2(dim)
    index.add(np.array(vectors))
    return index, vectors, chunks

def search_similar_chunks(query, index, chunks, vectors, top_k=3):
    q_vec = embed_texts([query])[0]
    D, I = index.search(np.array([q_vec]), top_k)
    return [chunks[i] for i in I[0]]

def call_ai(prompt):
    response = client.chat.completions.create(
        model=GPT_DEPLOYMENT,
        messages=[
            {"role": "system", "content": "Tu es un expert CIR."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.4,
        max_tokens=1500
    )
    return response.choices[0].message.content

def generate_section_with_rag(titre, question, index, chunks, vectors):
    context = "\n".join(search_similar_chunks(question, index, chunks, vectors))
    prompt = f"""Tu es un consultant CIR. Tu dois rédiger la section : "{titre}".

Voici le contexte extrait des documents client :
\"\"\"
{context}
\"\"\"

Rédige cette section de façon claire et structurée."""
    return call_ai(prompt)

def generer_etat_art(articles):
    refs = [f"- {a['title']} ({a['year']}) — {a['authors']} — {a['url']}"
            for a in articles if a.get("selected")]
    return "## État de l’art scientifique\n\n" + ("\n".join(refs) if refs else "Aucun article retenu.")

def remplir_doc(template_path, output_path, sections):
    doc = Document(template_path)
    for i, para in enumerate(doc.paragraphs):
        titre = para.text.strip()
        if titre in sections:
            if i + 1 < len(doc.paragraphs):
                doc.paragraphs[i + 1].text = sections[titre]
            else:
                doc.add_paragraph(sections[titre])
    doc.save(output_path)

# ─────── UI STREAMLIT ────────────────
st.set_page_config(page_title="Assistant CIR", layout="wide")
st.title("🧠 Assistant CIR – RAG + Articles scientifiques")

projet_name = st.text_input("Nom du projet")
uploaded_files = st.file_uploader("📎 Documents client (PDF ou DOCX)", type=["pdf", "docx"], accept_multiple_files=True)

if uploaded_files:
    # Indexation des documents
    full_text = ""
    for f in uploaded_files:
        st.write(f"📄 Lecture : {f.name}")
        full_text += "\n" + extract_text(f)

    chunks = chunk_text(full_text)
    index, vectors, raw_chunks = build_faiss_index(chunks)
    st.success(f"✅ {len(chunks)} morceaux indexés avec succès.")

    # Simuler articles scientifiques (à remplacer par recherche CrossRef réelle)
    if "articles" not in st.session_state:
        st.session_state["articles"] = [
            {"title": "AI in healthcare", "year": 2024, "authors": "Smith et al.", "url": "https://doi.org/fake1", "selected": True},
            {"title": "Deep learning for R&D", "year": 2023, "authors": "Jones et al.", "url": "https://doi.org/fake2", "selected": True},
            {"title": "Old paper", "year": 2011, "authors": "Doe", "url": "https://doi.org/old", "selected": False},
        ]

    st.subheader("📚 Articles scientifiques proposés")
    for i, a in enumerate(st.session_state["articles"]):
        key = f"article_{i}"
        st.session_state["articles"][i]["selected"] = st.checkbox(
            f"{a['title']} ({a['year']}) — {a['authors']}", value=a["selected"], key=key
        )

    if st.button("✨ Générer le dossier CIR"):
        with st.spinner("✍️ Rédaction en cours…"):
            sections = {
                "Contexte de l’opération de R&D": generate_section_with_rag("Contexte", "Quel est le contexte scientifique et industriel du projet ?", index, raw_chunks, vectors),
                "Indicateurs de R&D": generate_section_with_rag("Indicateurs", "Quels sont les indicateurs R&D liés au projet ?", index, raw_chunks, vectors),
                "Objet de l’opération de R&D": generate_section_with_rag("Objectifs", "Quels sont les objectifs techniques et les verrous scientifiques ?", index, raw_chunks, vectors),
                "Description de la démarche suivie et des travaux réalisés": generate_section_with_rag("Travaux", "Quels travaux ont été menés dans le projet ?", index, raw_chunks, vectors),
                "Contribution scientifique, technique ou technologique": generate_section_with_rag("Contribution", "Quelle est la contribution scientifique ou technique du projet ?", index, raw_chunks, vectors),
                "Partenariat scientifique et recherche confiée": generate_section_with_rag("Partenariat", "Quels sont les partenariats ou sous-traitances impliqués ?", index, raw_chunks, vectors),
                "État de l’art scientifique": generer_etat_art(st.session_state["articles"])
            }

            template_path = "./Doc/CLIENT_CIR.docx"
            output_path = f"./Doc/Dossier_CIR_{projet_name.replace(' ', '_')}.docx"
            remplir_doc(template_path, output_path, sections)

            st.success("✅ Dossier généré avec succès !")
            with open(output_path, "rb") as f:
                st.download_button("📥 Télécharger le dossier Word", f, file_name=os.path.basename(output_path))
