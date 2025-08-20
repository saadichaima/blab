# Core/writer.py
from __future__ import annotations

from io import BytesIO
from typing import Iterable, Optional, Dict
import re
import unicodedata

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

# ───────────────────────── Helpers généraux ─────────────────────────

def _normalize(s: str) -> str:
    """Nettoie espaces multiples / NBSP pour des comparaisons robustes."""
    return " ".join((s or "").replace("\xa0", " ").split()).strip()

def _clear_container(hf):
    """Supprime paragraphes et tableaux d'un header/footer pour réécrire proprement."""
    for tbl in list(hf.tables):
        tbl._element.getparent().remove(tbl._element)
    for p in list(hf.paragraphs):
        p._element.getparent().remove(p._element)

def _delete_paragraph(p):
    p._element.getparent().remove(p._element)

def _ensure_paragraph_after(paragraph):
    """S'assure qu'il y a un paragraphe juste après 'paragraph' et le retourne."""
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, paragraph._parent)

# Remplacement textuel qui préserve la mise en forme (travaille par runs)
def _replace_in_runs(paragraph, target: str, replacement: str) -> bool:
    """
    Remplace target par replacement dans un paragraphe.
    Si le mot est coupé en plusieurs runs, on bascule sur un remplacement
    'plein paragraphe' (perte locale du style dans ce paragraphe).
    """
    done = False
    for r in paragraph.runs:
        if target in r.text:
            r.text = r.text.replace(target, replacement)
            done = True
    if (not done) and target in paragraph.text:
        paragraph.text = paragraph.text.replace(target, replacement)
        done = True
    return done

def _replace_in_paragraphs(paragraphs: Iterable, target: str, replacement: str) -> int:
    count = 0
    for p in paragraphs:
        if _replace_in_runs(p, target, replacement):
            count += 1
    return count

def _replace_in_tables(tables, target: str, replacement: str) -> int:
    count = 0
    for tbl in tables:
        for row in tbl.rows:
            for cell in row.cells:
                count += _replace_in_paragraphs(cell.paragraphs, target, replacement)
    return count

# ───────────────────────── Helpers titres ─────────────────────────

def _strip_accents(s: str) -> str:
    return "".join(c for c in unicodedata.normalize("NFKD", s or "") if not unicodedata.combining(c))

def _normalize_title(s: str) -> str:
    """
    Normalise un libellé de titre pour comparaison tolérante:
    - remplace NBSP, compacte espaces
    - supprime numérotation en tête: '1.', '2)', 'I -', '1.1', etc.
    - passe en lower, supprime accents
    - remplace apostrophes typographiques par simples
    """
    s = (s or "").replace("\xa0", " ").strip()
    s = re.sub(r"^\s*[\dIVXLCM]+(?:\.\d+)*[\.\)\-–—]?\s+", "", s)  # 1., 2), I -, 1.1, etc.
    s = s.replace("’", "'")
    s = " ".join(s.split())
    s = _strip_accents(s).lower()
    return s

def _is_heading(p) -> bool:
    """Vrai si le paragraphe a un style de type Titre/Heading."""
    if not p.style:
        return False
    name = (p.style.name or "").lower()
    return name.startswith("heading") or name.startswith("titre")

# ───────────────────────── Remplacements globaux (CLIENT / LOGO / 20XX) ─────────────────────────

def _replace_all_text_everywhere(doc: Document, target: str, replacement: str) -> int:
    """
    Remplace target par replacement dans tout le document :
    corps, tableaux, en-têtes et pieds (y compris 'première page').
    Retourne le nombre de paragraphes modifiés.
    """
    changed = 0
    # Corps
    changed += _replace_in_paragraphs(doc.paragraphs, target, replacement)
    changed += _replace_in_tables(doc.tables, target, replacement)
    # En-têtes / pieds
    for sec in doc.sections:
        for header in [sec.header, sec.first_page_header]:
            changed += _replace_in_paragraphs(header.paragraphs, target, replacement)
            changed += _replace_in_tables(header.tables, target, replacement)
        for footer in [sec.footer, sec.first_page_footer]:
            changed += _replace_in_paragraphs(footer.paragraphs, target, replacement)
            changed += _replace_in_tables(footer.tables, target, replacement)
    return changed

def _insert_logo_everywhere(doc: Document, logo_bytes: Optional[bytes], width_inches: float = 2.0) -> int:
    """
    Remplace chaque paragraphe qui contient 'LOGO' par l'image donnée.
    (corps, tableaux, en-têtes et pieds)
    Retourne le nombre d'emplacements remplacés.
    """
    if not logo_bytes:
        return 0

    def replace_logo_in_paragraphs(paragraphs):
        cnt = 0
        for p in paragraphs:
            if "LOGO" in p.text:
                p.text = ""
                run = p.add_run()
                run.add_picture(BytesIO(logo_bytes), width=Inches(width_inches))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cnt += 1
        return cnt

    def replace_logo_in_tables(tables):
        cnt = 0
        for tbl in tables:
            for row in tbl.rows:
                for cell in row.cells:
                    cnt += replace_logo_in_paragraphs(cell.paragraphs)
        return cnt

    changed = 0
    # Corps
    changed += replace_logo_in_paragraphs(doc.paragraphs)
    changed += replace_logo_in_tables(doc.tables)
    # En-têtes / pieds
    for sec in doc.sections:
        for header in [sec.header, sec.first_page_header]:
            changed += replace_logo_in_paragraphs(header.paragraphs)
            changed += replace_logo_in_tables(header.tables)
        for footer in [sec.footer, sec.first_page_footer]:
            changed += replace_logo_in_paragraphs(footer.paragraphs)
            changed += replace_logo_in_tables(footer.tables)
    return changed

def _replace_all_20XX(doc: Document, year: int) -> int:
    """Remplace toutes les occurrences exactes de '20XX' par l'année partout (corps, titres, en-têtes/pieds, tableaux)."""
    return _replace_all_text_everywhere(doc, "20XX", str(year))

def _replace_all_CLIENT(doc: Document, client: str) -> int:
    """Remplace toutes les occurrences de 'CLIENT' partout."""
    return _replace_all_text_everywhere(doc, "CLIENT", client)

# ───────────────────────── Pieds de page ─────────────────────────

def _set_footer_all_sections(doc: Document, client: str, year: int):
    """
    Pieds de page pour toutes les sections :
      - Ligne 1 : gauche, gras -> 'CLIENT – CIR YEAR'
      - Ligne 2 : centrée      -> 'Mémoire CIR YEAR'
    Couvre aussi le cas 'Première page différente'.
    """
    top = f"{client} – CIR {year}"
    bottom = f"Mémoire CIR {year}"

    for sec in doc.sections:
        sec.footer.is_linked_to_previous = False
        sec.first_page_footer.is_linked_to_previous = False

        for foot in (sec.first_page_footer, sec.footer):
            _clear_container(foot)

            p1 = foot.add_paragraph()
            p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r1 = p1.add_run(top)
            r1.bold = True

            p2 = foot.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.add_run(bottom)

# ───────────────────────── Branding 1ʳᵉ page ─────────────────────────

def apply_branding_first_page(
    template_path: str,
    output_path: str,
    client: str,
    year: int,
    logo_bytes: Optional[bytes] = None
):
    """
    Personnalise le modèle AVANT remplissage :
      - 'CLIENT' -> nom de la société (partout)
      - 'LOGO'   -> image fournie à la place de chaque occurence (corps/en-têtes/pieds/tableaux)
      - '20XX'   -> année donnée (partout)
      - Pieds de page formatés sur toutes les pages.
    Sauvegarde la copie personnalisée dans output_path.
    """
    doc = Document(template_path)

    # Remplacements globaux (corps + en-têtes/pieds + tableaux)
    _replace_all_CLIENT(doc, client)
    _insert_logo_everywhere(doc, logo_bytes)
    _replace_all_20XX(doc, year)

    # Puis on pose des pieds de page cohérents
    _set_footer_all_sections(doc, client, year)

    doc.save(output_path)

# ───────────────────────── Contenu IA (état de l’art) ─────────────────────────

def generer_etat_art(articles):
    """Formate la section 'État de l’art scientifique' depuis les articles sélectionnés."""
    refs = [
        f"- {a['title']} ({a['year']}) — {a['authors']} — {a.get('url','')}"
        for a in articles if a.get("selected")
    ]
    return "## État de l’art scientifique\n\n" + ("\n".join(refs) if refs else "Aucun article retenu.")

# ───────────────────────── Remplissage des sections (matching strict) ─────────────────────────

def remplir_doc(template_path: str, output_path: str, sections: Dict[str, str]):
    """
    Remplit le document en remplaçant, pour chaque Titre (Heading/Titre 1..9) qui
    correspond à une clé de `sections`, tout le contenu jusqu'au prochain titre.

    Matching STRICT :
      - égalité après normalisation (numérotation, NBSP, accents, casse, apostrophes)
      - pas de correspondance partielle → n'affecte pas 'PRÉSENTATION DE L’ENTREPRISE'
        quand on cible 'L’entreprise'.

    ⚠️ Les notes de bas de page sont gérées APRÈS sauvegarde
       par Core.footnotes.add_smart_footnotes(output_path).
    """
    doc = Document(template_path)

    # Prépare un mapping titres normalisés -> (titre d'origine, contenu)
    targets = { _normalize_title(t): (t, content) for t, content in sections.items() }
    remaining = set(targets.keys())

    i = 0
    while i < len(doc.paragraphs):
        p = doc.paragraphs[i]

        if not _is_heading(p):
            i += 1
            continue

        norm = _normalize_title(p.text)

        # correspondance stricte uniquement (après normalisation)
        key = norm if norm in targets else None
        if key is None:
            i += 1
            continue

        # On a un titre ciblé
        original_title, content = targets[key]
        remaining.discard(key)

        # Trouver l'index du prochain heading
        j = i + 1
        while j < len(doc.paragraphs) and not _is_heading(doc.paragraphs[j]):
            j += 1
        # Bloc à remplacer = (i+1) .. (j-1)

        # Cas 1 : aucun paragraphe après le titre (ou directement un autre titre)
        if i + 1 >= len(doc.paragraphs) or j == i + 1:
            after = _ensure_paragraph_after(p)
            after.text = content or ""
            i += 2  # on saute le titre + le para inséré
            continue

        # Cas 2 : il y a du contenu → on remplace le 1er et on supprime le reste
        first_content_para = doc.paragraphs[i + 1]
        first_content_para.text = content or ""

        end = j if j <= len(doc.paragraphs) else len(doc.paragraphs)
        # Supprimer les suivants (de la fin vers le début)
        for k in range(end - 1, i + 1, -1):
            _delete_paragraph(doc.paragraphs[k])

        i += 2  # titre + 1 paragraphe de contenu

    # Les titres demandés mais absents du modèle → ajout en fin
    if remaining:
        for k in remaining:
            original_title, content = targets[k]
            doc.add_heading(original_title, level=2)
            doc.add_paragraph(content or "")

    doc.save(output_path)
