import os
import uuid
import fitz  # PyMuPDF
from docx import Document
from dotenv import load_dotenv
from openai import AzureOpenAI
import chromadb
from langchain_openai import AzureOpenAIEmbeddings




# Load env vars
load_dotenv()

# Azure OpenAI
client = AzureOpenAI(
    api_key=os.getenv("AZURE_OPENAI_KEY"),
    api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
    azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
)

deployment_name = os.getenv("AZURE_OPENAI_DEPLOYMENT")
embedding_deployment = os.getenv("AZURE_OPENAI_EMBEDDING_DEPLOYMENT")

# Init Chroma
chroma_client = chromadb.PersistentClient(path="./rag_db")
collection = chroma_client.get_or_create_collection(name="client_chunks")

# 1. Lire tous les fichiers texte (PDF, Word)
def extract_text_from_folder(folder_path):
    full_text = ""
    for file in os.listdir(folder_path):
        path = os.path.join(folder_path, file)
        if file.endswith(".pdf"):
            doc = fitz.open(path)
            for page in doc:
                full_text += page.get_text()
            doc.close()
        elif file.endswith(".docx"):
            doc = Document(path)
            for para in doc.paragraphs:
                full_text += para.text + "\n"
        elif file.endswith(".txt"):
            with open(path, "r", encoding="utf-8") as f:
                full_text += f.read()
    return full_text

# 2. Chunker le texte
def chunk_text(text, max_length=1000, overlap=200):
    chunks = []
    start = 0
    while start < len(text):
        end = start + max_length
        chunks.append(text[start:end])
        start += max_length - overlap
    return chunks

# 3. Indexer les chunks dans ChromaDB
def index_chunks(chunks):
    embedding_model = AzureOpenAIEmbeddings(
    azure_deployment=os.getenv("AZURE_OPENAI_EMBEDDING_DEPLOYMENT"),
    openai_api_key=os.getenv("AZURE_OPENAI_KEY"),
    azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),  
    openai_api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
    openai_api_type="azure"
)

    for chunk in chunks:
        vector = embedding_model.embed_query(chunk)
        collection.add(
            ids=[str(uuid.uuid4())],
            documents=[chunk],
            embeddings=[vector]
        )

# 4. Récupérer les chunks pertinents
def retrieve_relevant_chunks(question, top_k=5):
    embedding_model = AzureOpenAIEmbeddings(
    azure_deployment=os.getenv("AZURE_OPENAI_EMBEDDING_DEPLOYMENT"),
    openai_api_key=os.getenv("AZURE_OPENAI_KEY"),
    azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),  
    openai_api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
    openai_api_type="azure"
)

    query_vector = embedding_model.embed_query(question)
    results = collection.query(
        query_embeddings=[query_vector],
        n_results=top_k
    )
    return results["documents"][0]

# 5. Génération du rapport basé sur les chunks
def generate_rag_report(question):
    chunks = retrieve_relevant_chunks(question)
    context = "\n\n".join(chunks)
    prompt = f"""
    À partir des extraits suivants des documents clients :
    {context}

    Réponds à la question suivante :
    {question}
    """

    response = client.chat.completions.create(
        model=deployment_name,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.4,
        max_tokens=1000
    )

    return response.choices[0].message.content

# 6. Sauvegarder le rapport dans un docx
def save_report_to_word(report_text, output_path):
    doc = Document()
    doc.add_heading("Rapport RAG Généré", 0)
    for line in report_text.split("\n"):
        doc.add_paragraph(line.strip())
    doc.save(output_path)

# Main
if __name__ == "__main__":
    folder = "Doc"
    print("🔍 Lecture des documents...")
    text = extract_text_from_folder(folder)
    chunks = chunk_text(text)

    print("📦 Indexation des chunks...")
    index_chunks(chunks)

    question = "Synthétise les objectifs et les performances visées dans le projet"
    print(f"❓ Question posée : {question}")
    report = generate_rag_report(question)

    print("\n📝 Rapport généré :\n")
    print(report)

    save_report_to_word(report, "rapport_rag.docx")
    print("✅ Rapport sauvegardé sous rapport_rag.docx")
